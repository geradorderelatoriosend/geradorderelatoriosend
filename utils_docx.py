import os
from datetime import date
from docx.shared import Inches, Cm
from docx import Document


def format_date_br(d: date | None) -> str:
    if not d:
        return ""
    return d.strftime("%d/%m/%Y")


def date_por_extenso(d: date | None) -> str:
    if not d:
        return ""
    meses = [
        "janeiro", "fevereiro", "março", "abril", "maio", "junho",
        "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"
    ]
    return f"{d.day} de {meses[d.month - 1]} de {d.year}"


def replace_placeholder_in_paragraph(paragraph, key: str, value: str):
    """
    Substitui o placeholder apenas dentro dos runs,
    para manter a formatação original (fonte, tamanho, cor, sombra etc.).
    """
    if not key or key not in paragraph.text:
        return

    for run in paragraph.runs:
        if key in run.text:
            run.text = run.text.replace(key, value)

    # ❌ NÃO reescrevemos paragraph.text inteiro,
    # para não perder a formatação original.


def _replace_in_paragraphs_collection(paragraphs, mapping: dict):
    """Aplica a substituição em uma coleção de parágrafos (corpo, cabeçalho, rodapé, células etc.)."""
    for p in paragraphs:
        for key, value in mapping.items():
            if key in p.text:
                replace_placeholder_in_paragraph(p, key, value)


def _replace_in_tables_collection(tables, mapping: dict):
    """Aplica a substituição em uma coleção de tabelas (corpo, cabeçalhos, rodapés)."""
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_in_paragraphs_collection(cell.paragraphs, mapping)


def replace_placeholders(document: Document, mapping: dict):
    """
    Substitui placeholders em todo o documento:
      - corpo do documento (parágrafos e tabelas)
      - cabeçalhos e rodapés
      - shapes / caixas de texto (via XML)
    mapping: dict do tipo {"{{PLACEHOLDER}}": "valor"}.
    """
    # ---------- Corpo principal ----------
    _replace_in_paragraphs_collection(document.paragraphs, mapping)
    _replace_in_tables_collection(document.tables, mapping)

    # ---------- Cabeçalhos e rodapés ----------
    for section in document.sections:
        header = section.header
        footer = section.footer

        if header is not None:
            _replace_in_paragraphs_collection(header.paragraphs, mapping)
            _replace_in_tables_collection(header.tables, mapping)

        if footer is not None:
            _replace_in_paragraphs_collection(footer.paragraphs, mapping)
            _replace_in_tables_collection(footer.tables, mapping)

    # ---------- Shapes / caixas de texto / outros w:t no XML ----------
    root = document.element
    for key, value in mapping.items():
        if not key:
            continue
        for node in root.iter():
            # w:t (texto) – o sufixo do tag é sempre '}t'
            if node.tag.endswith('}t') and node.text and key in node.text:
                node.text = node.text.replace(key, value)


def inserir_foto_por_placeholder(
    doc: Document,
    placeholder: str,
    image_path: str,
    width_inches: float | None = None,
    height_cm: float | None = 4.0,
):
    """
    Insere uma imagem no lugar de um placeholder dentro do documento.
    - Se height_cm for informado, usa altura fixa em cm (largura proporcional).
    - Caso contrário, usa width_inches em polegadas.
    """
    if not image_path or not os.path.exists(image_path):
        return

    # Parágrafos fora de tabela
    for p in doc.paragraphs:
        if placeholder in p.text:
            p.text = p.text.replace(placeholder, "")
            run = p.add_run()
            if height_cm is not None:
                run.add_picture(image_path, height=Cm(height_cm))
            elif width_inches is not None:
                run.add_picture(image_path, width=Inches(width_inches))
            else:
                run.add_picture(image_path)
            return

    # Parágrafos dentro de tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if placeholder in p.text:
                        p.text = p.text.replace(placeholder, "")
                        run = p.add_run()
                        if height_cm is not None:
                            run.add_picture(image_path, height=Cm(height_cm))
                        elif width_inches is not None:
                            run.add_picture(image_path, width=Inches(width_inches))
                        else:
                            run.add_picture(image_path)
                        return
